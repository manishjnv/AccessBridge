"""Generate AccessBridge Manual Test Cases Word Document"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

doc = Document()

# Page setup
for section in doc.sections:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ---- Title Page ----
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AccessBridge')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(123, 104, 238)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Manual Test Cases')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Version: 0.1.1\n').font.size = Pt(12)
meta.add_run('Wipro TopGear Ideathon 2026\n').font.size = Pt(12)
meta.add_run('Author: Manish Kumar\n').font.size = Pt(12)
meta.add_run('Date: April 2026').font.size = Pt(12)

doc.add_page_break()

# ---- Table of Contents placeholder ----
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Prerequisites & Setup',
    '2. Extension Installation (TC-001 to TC-003)',
    '3. Popup UI & Navigation (TC-004 to TC-009)',
    '4. Sensory Adaptations (TC-010 to TC-016)',
    '5. Cognitive Features (TC-017 to TC-023)',
    '6. Motor Assistor (TC-024 to TC-033)',
    '7. AI Engine (TC-034 to TC-039)',
    '8. Domain Connectors (TC-040 to TC-043)',
    '9. Side Panel (TC-044 to TC-047)',
    '10. Self-Update System (TC-048 to TC-050)',
    '11. Profile Management (TC-051 to TC-054)',
    '12. VPS & Landing Page (TC-055 to TC-058)',
    '13. Regression & Edge Cases (TC-059 to TC-063)',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ---- Helper to add test case table ----
tc_counter = [0]

def add_tc(tc_id, title, preconditions, steps, expected, priority='High', category=''):
    tc_counter[0] += 1
    doc.add_heading(f'{tc_id}: {title}', level=3)

    table = doc.add_table(rows=6, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(13.5)

    fields = [
        ('Test Case ID', tc_id),
        ('Category', category),
        ('Priority', priority),
        ('Preconditions', preconditions),
        ('Steps', steps),
        ('Expected Result', expected),
    ]

    for i, (label, value) in enumerate(fields):
        cell_label = table.rows[i].cells[0]
        cell_value = table.rows[i].cells[1]

        run = cell_label.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)

        cell_value.paragraphs[0].text = ''
        for line in value.split('\n'):
            run = cell_value.paragraphs[0].add_run(line + '\n')
            run.font.size = Pt(9)

    doc.add_paragraph()  # spacing


# ============================================================
# SECTION 1: PREREQUISITES
# ============================================================
doc.add_heading('1. Prerequisites & Setup', level=1)
p = doc.add_paragraph()
p.add_run('Environment Requirements:\n').bold = True
doc.add_paragraph('Google Chrome (v120+) with Developer Mode enabled', style='List Bullet')
doc.add_paragraph('AccessBridge extension dist/ folder (from pnpm build)', style='List Bullet')
doc.add_paragraph('Internet access to VPS: http://72.61.227.64:8300/', style='List Bullet')
doc.add_paragraph('Microphone access (for voice commands)', style='List Bullet')
doc.add_paragraph('Webcam access (for eye tracking)', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Test Data:\n').bold = True
doc.add_paragraph('Gmail account with at least one email', style='List Bullet')
doc.add_paragraph('Any banking website (e.g. SBI, HDFC online banking)', style='List Bullet')
doc.add_paragraph('Any insurance website (e.g. LIC, ICICI Prudential)', style='List Bullet')
doc.add_paragraph()

# ============================================================
# SECTION 2: EXTENSION INSTALLATION
# ============================================================
doc.add_heading('2. Extension Installation', level=1)

add_tc('TC-001', 'Sideload Extension in Chrome',
    'Chrome browser open, Developer Mode enabled in chrome://extensions',
    '1. Navigate to chrome://extensions\n'
    '2. Enable "Developer mode" toggle (top right)\n'
    '3. Click "Load unpacked"\n'
    '4. Select the packages/extension/dist/ folder\n'
    '5. Observe the extension card appears',
    'Extension loads without errors.\n'
    'Extension card shows: Name "AccessBridge", Version "0.1.1", < 1 MB size.\n'
    'Service worker shows as active (not "(Inactive)").\n'
    'Extension icon appears in toolbar.',
    category='Installation')

add_tc('TC-002', 'Verify Manifest & Permissions',
    'Extension loaded per TC-001',
    '1. On chrome://extensions, click "Details" on AccessBridge\n'
    '2. Check Permissions section\n'
    '3. Check Site access setting\n'
    '4. Verify "Inspect views: service worker" link exists',
    'Permissions include: activeTab, storage, offscreen, sidePanel.\n'
    'Site access: "On all sites" (host_permissions: <all_urls>).\n'
    'Service worker link is visible and clickable.',
    category='Installation')

add_tc('TC-003', 'Service Worker Console Clean',
    'Extension loaded per TC-001',
    '1. Click "service worker" link under Inspect views\n'
    '2. DevTools opens — check Console tab\n'
    '3. Look for error messages',
    'Console shows:\n'
    '  - "AccessBridge service worker initialized"\n'
    '  - "AccessBridge installed (reason: install, version: 0.1.1)"\n'
    'No errors (red messages).\n'
    'No warnings related to AccessBridge.',
    category='Installation')

# ============================================================
# SECTION 3: POPUP UI
# ============================================================
doc.add_heading('3. Popup UI & Navigation', level=1)

add_tc('TC-004', 'Popup Opens and Displays Header',
    'Extension loaded, icon pinned to toolbar',
    '1. Click the AccessBridge icon in Chrome toolbar\n'
    '2. Observe the popup window',
    'Popup opens (360px wide).\n'
    'Header shows "AccessBridge" in accent color and "v0.1.1".\n'
    'Master toggle (On/Off) is visible and set to On.\n'
    'No update banner shown (since version matches server).',
    category='Popup UI')

add_tc('TC-005', 'Tab Navigation Works',
    'Popup open per TC-004',
    '1. Click "Sensory" tab\n'
    '2. Click "Cognitive" tab\n'
    '3. Click "Motor" tab\n'
    '4. Click "Settings" tab\n'
    '5. Click "Overview" tab to return',
    'Each tab switches content correctly.\n'
    'Active tab is visually highlighted.\n'
    'No flickering or layout shift during transitions.',
    category='Popup UI')

add_tc('TC-006', 'Overview Tab — Struggle Score Display',
    'Popup open, Overview tab selected',
    '1. Observe the struggle score display\n'
    '2. Note the score value and color\n'
    '3. Check Active Adaptations count\n'
    '4. Check Struggle Level label',
    'Struggle Score shows "0" (green color) initially.\n'
    'Active Adaptations: 0.\n'
    'Struggle Level: "Low".\n'
    'Score updates every 3 seconds (polling).',
    category='Popup UI')

add_tc('TC-007', 'Overview Tab — Score Updates on User Activity',
    'Popup open, any webpage open in active tab',
    '1. Open a webpage (e.g. google.com)\n'
    '2. Rapidly scroll up and down for 10 seconds\n'
    '3. Click random areas rapidly\n'
    '4. Open popup and check Overview tab',
    'Struggle Score increases from 0 (reflects detected struggle).\n'
    'Score color changes: green (<30), yellow (30-60), red (>60).\n'
    'Active Adaptations may increase if auto-adaptations triggered.',
    category='Popup UI')

add_tc('TC-008', 'Master Toggle — Disable All',
    'Popup open, some features enabled',
    '1. Toggle the master switch to Off\n'
    '2. Navigate to any webpage\n'
    '3. Verify no AccessBridge UI elements appear on page',
    'All active adaptations are reverted.\n'
    'No AccessBridge overlays, cursors, or UI injected on pages.\n'
    'Popup still opens but features are inactive.',
    category='Popup UI')

add_tc('TC-009', 'Master Toggle — Re-enable',
    'Master toggle is Off from TC-008',
    '1. Toggle the master switch back to On\n'
    '2. Navigate to a webpage\n'
    '3. Enable a feature (e.g. Focus Mode)',
    'Features can be activated again.\n'
    'Enabled features apply to the page normally.',
    category='Popup UI')

# ============================================================
# SECTION 4: SENSORY ADAPTATIONS
# ============================================================
doc.add_heading('4. Sensory Adaptations', level=1)

add_tc('TC-010', 'Font Scale Slider',
    'Popup open, Sensory tab, any webpage in active tab',
    '1. Open a text-heavy webpage (e.g. Wikipedia article)\n'
    '2. Open popup > Sensory tab\n'
    '3. Drag Font Scale slider from 1.0x to 1.5x\n'
    '4. Observe the webpage text\n'
    '5. Drag to 2.0x (maximum)',
    'Page text scales proportionally.\n'
    'At 1.5x, text is noticeably larger.\n'
    'At 2.0x, text is double the original size.\n'
    'Layout remains usable (no overlapping text).',
    category='Sensory')

add_tc('TC-011', 'Contrast Slider',
    'Popup open, Sensory tab, any webpage',
    '1. Drag Contrast slider to 1.5x\n'
    '2. Observe page colors\n'
    '3. Drag to 2.0x (maximum)\n'
    '4. Return to 1.0x',
    'Contrast increases visibly at 1.5x.\n'
    'At 2.0x, high contrast mode — dark backgrounds darker, light text brighter.\n'
    'Returning to 1.0x restores original appearance.',
    category='Sensory')

add_tc('TC-012', 'Line Height Adjustment',
    'Popup open, Sensory tab, text-heavy page',
    '1. Set Line Height slider to 2.0\n'
    '2. Observe paragraph spacing\n'
    '3. Set to 3.0 (maximum)',
    'Line spacing between text lines increases.\n'
    'Text becomes easier to read with more vertical space.\n'
    'Page scrolls longer but text is more spread out.',
    category='Sensory')

add_tc('TC-013', 'Letter Spacing Adjustment',
    'Popup open, Sensory tab, text-heavy page',
    '1. Set Letter Spacing slider to 2px\n'
    '2. Observe character spacing\n'
    '3. Set to 5px (maximum)',
    'Characters have visible gaps between them.\n'
    'At 5px, spacing is very wide but characters remain readable.\n'
    'Helps users with dyslexia distinguish letters.',
    category='Sensory')

add_tc('TC-014', 'Color Correction — Protanopia',
    'Popup open, Sensory tab, colorful webpage',
    '1. Open a page with red/green elements (e.g. a chart)\n'
    '2. Select Color Correction > "Protanopia (Red-blind)"\n'
    '3. Observe color changes',
    'Page colors shift to compensate for red-blind vision.\n'
    'Red elements become more distinguishable.\n'
    'Filter applies as a CSS overlay.',
    category='Sensory')

add_tc('TC-015', 'Reduced Motion Toggle',
    'Popup open, Sensory tab, page with animations',
    '1. Open a page with CSS animations or transitions\n'
    '2. Toggle "Reduced Motion" On\n'
    '3. Observe animations',
    'CSS animations and transitions are disabled or reduced.\n'
    'prefers-reduced-motion media query is effectively applied.\n'
    'Page content remains accessible but without movement.',
    category='Sensory')

add_tc('TC-016', 'High Contrast Toggle',
    'Popup open, Sensory tab',
    '1. Toggle "High Contrast" On\n'
    '2. Observe page appearance',
    'Page switches to high-contrast color scheme.\n'
    'Text becomes sharply contrasted against background.\n'
    'Useful for low-vision users.',
    category='Sensory')

# ============================================================
# SECTION 5: COGNITIVE FEATURES
# ============================================================
doc.add_heading('5. Cognitive Features', level=1)

add_tc('TC-017', 'Focus Mode',
    'Popup open, Cognitive tab, any webpage',
    '1. Toggle "Focus Mode" On\n'
    '2. Observe the webpage\n'
    '3. Move mouse to different areas\n'
    '4. Toggle Off to deactivate',
    'A spotlight effect appears around the cursor area.\n'
    'Surrounding content is dimmed/darkened.\n'
    'Only the area near the cursor is fully visible.\n'
    'Toggling Off removes the spotlight overlay.',
    category='Cognitive')

add_tc('TC-018', 'Reading Mode',
    'Popup open, Cognitive tab, article/blog page',
    '1. Toggle "Reading Mode" On\n'
    '2. Observe the page',
    'A reading guide line/ruler appears on the page.\n'
    'Helps track which line the user is reading.\n'
    'Guide follows vertical mouse position.',
    category='Cognitive')

add_tc('TC-019', 'Text Simplification — Mild',
    'Popup open, Cognitive tab, text-heavy page',
    '1. Set Text Simplification dropdown to "Mild"\n'
    '2. Observe page text',
    'Complex words replaced with simpler alternatives.\n'
    'e.g. "utilize" → "use", "facilitate" → "help".\n'
    'Page remains readable with natural flow.',
    category='Cognitive')

add_tc('TC-020', 'Text Simplification — Strong',
    'Popup open, Cognitive tab, text-heavy page',
    '1. Set Text Simplification dropdown to "Strong"\n'
    '2. Observe page text',
    'Complex words replaced AND long sentences broken up.\n'
    'Parenthetical asides removed.\n'
    'Sentences split at conjunctions (and, but, because).\n'
    'More aggressive simplification than Mild.',
    category='Cognitive')

add_tc('TC-021', 'Auto Summarize Toggle',
    'Popup open, Cognitive tab, article page',
    '1. Toggle "Auto Summarize" On\n'
    '2. Observe the page',
    'A summary panel appears with bullet points of page content.\n'
    'Panel shows: tier (local), latency, bullet points.\n'
    'Buttons: Simplify, Read Aloud, Copy.',
    category='Cognitive')

add_tc('TC-022', 'Distraction Shield',
    'Popup open, Cognitive tab, page with ads/banners/sidebars',
    '1. Toggle "Distraction Shield" On\n'
    '2. Observe the page',
    'Sidebar content, ads, banners, and non-essential elements hidden or dimmed.\n'
    'Main content remains visible and prominent.\n'
    'Page is cleaner and less visually noisy.',
    category='Cognitive')

add_tc('TC-023', 'Notification Level Setting',
    'Popup open, Cognitive tab',
    '1. Change Notification Level to "Important Only"\n'
    '2. Change to "Critical Only"\n'
    '3. Change to "None"',
    'Dropdown saves selection to profile.\n'
    'Setting persists after closing and reopening popup.\n'
    '(Notification filtering applies to supported web apps.)',
    category='Cognitive')

# ============================================================
# SECTION 6: MOTOR ASSISTOR
# ============================================================
doc.add_heading('6. Motor Assistor Features', level=1)

add_tc('TC-024', 'Voice Navigation — Enable',
    'Popup open, Motor tab, microphone permission granted',
    '1. Toggle "Voice Navigation" On\n'
    '2. Browser may prompt for microphone — click Allow\n'
    '3. Observe speech recognition indicator',
    'Voice command system activates.\n'
    'Web Speech API starts listening.\n'
    'No errors in console.',
    category='Motor')

add_tc('TC-025', 'Voice Command — Scroll',
    'Voice Navigation enabled per TC-024, any long webpage',
    '1. Say "scroll down"\n'
    '2. Say "scroll up"\n'
    '3. Say "go to top"\n'
    '4. Say "go to bottom"',
    'Page scrolls down ~300px on "scroll down".\n'
    'Page scrolls up ~300px on "scroll up".\n'
    '"go to top" scrolls to page top.\n'
    '"go to bottom" scrolls to page bottom.\n'
    'All with smooth scrolling animation.',
    category='Motor')

add_tc('TC-026', 'Voice Command — Navigation',
    'Voice Navigation enabled, multiple tabs open',
    '1. Say "next tab"\n'
    '2. Say "previous tab"\n'
    '3. Say "new tab"\n'
    '4. Say "close tab"',
    '"next tab" switches to next browser tab.\n'
    '"previous tab" switches to previous tab.\n'
    '"new tab" opens a new empty tab.\n'
    '"close tab" closes the current tab.',
    category='Motor')

add_tc('TC-027', 'Voice Command — Click & Type',
    'Voice Navigation enabled, page with links/buttons',
    '1. Say "click [button text]" (e.g. "click Sign In")\n'
    '2. Focus a text input\n'
    '3. Say "type hello world"',
    '"click" finds and clicks the element matching the text.\n'
    '"type" inserts text into the focused input field.\n'
    'Input event fires (form validation works).',
    category='Motor')

add_tc('TC-028', 'Voice Command — Zoom',
    'Voice Navigation enabled',
    '1. Say "zoom in"\n'
    '2. Observe page zoom\n'
    '3. Say "zoom out"',
    '"zoom in" increases page zoom by 10%.\n'
    '"zoom out" decreases page zoom by 10%.\n'
    'Minimum zoom is 50%.',
    category='Motor')

add_tc('TC-029', 'Hindi Voice Commands',
    'Voice Navigation enabled, language set to Hindi in Settings',
    '1. Go to Settings tab, set Language to "Hindi"\n'
    '2. Close and reopen popup, enable Voice Navigation\n'
    '3. Say "neeche scroll karo" (scroll down)\n'
    '4. Say "upar scroll karo" (scroll up)\n'
    '5. Say "naya tab" (new tab)',
    'Hindi commands are recognized via Web Speech API with lang=hi-IN.\n'
    'Commands map to same actions as English equivalents.\n'
    '25+ Hindi command mappings supported.',
    category='Motor', priority='Medium')

add_tc('TC-030', 'Eye Tracking — Enable',
    'Popup open, Motor tab, webcam available',
    '1. Toggle "Eye Tracking" On\n'
    '2. Browser prompts for webcam — click Allow\n'
    '3. Observe the page',
    'Webcam preview appears in top-right corner.\n'
    'Label shows "Eye Tracker (FaceDetector)" if API available,\n'
    'or "Eye Tracker (fallback)" otherwise.\n'
    'Gaze cursor (semi-transparent circle) appears on page.\n'
    'Cursor moves roughly following head/face position.',
    category='Motor')

add_tc('TC-031', 'Dwell Click — Enable & Test',
    'Popup open, Motor tab, any page with buttons/links',
    '1. Toggle "Dwell Click" On\n'
    '2. Move mouse cursor over a button\n'
    '3. Hold still for ~1 second\n'
    '4. Observe the click indicator and action',
    'Radial SVG progress circle appears around cursor.\n'
    'After dwell delay (default ~800ms), auto-click fires.\n'
    'Visual pulse feedback on click.\n'
    'Button/link activates as if clicked.',
    category='Motor')

add_tc('TC-032', 'Keyboard-Only Mode',
    'Popup open, Motor tab, any complex webpage',
    '1. Toggle "Keyboard-Only Mode" On\n'
    '2. Press Tab key to navigate\n'
    '3. Press "?" key\n'
    '4. Press arrow keys in a group of elements\n'
    '5. Press Escape',
    'Skip links appear (Jump to Main, Jump to Nav, Jump to Footer).\n'
    'Enhanced focus ring visible on focused elements.\n'
    '"?" shows shortcuts overlay.\n'
    'Tab adds tabindex to clickable elements missing it.\n'
    'Arrow keys navigate within element groups.\n'
    'Escape deselects current focus.',
    category='Motor')

add_tc('TC-033', 'Predictive Input',
    'Popup open, Motor tab, page with text input',
    '1. Toggle "Predictive Input" On\n'
    '2. Click into a text input field\n'
    '3. Start typing "th"\n'
    '4. Observe suggestion panel\n'
    '5. Press Tab or Alt+1 to accept top suggestion',
    'Floating suggestion panel appears below input.\n'
    'Shows up to 5 word predictions (e.g. "the", "that", "this").\n'
    'Alt+1 through Alt+5 or Tab accepts a suggestion.\n'
    'Suggestions update as you type (80ms debounce).\n'
    'Form field intelligence: detects email/phone/address fields.',
    category='Motor')

# ============================================================
# SECTION 7: AI ENGINE
# ============================================================
doc.add_heading('7. AI Engine Features', level=1)

add_tc('TC-034', 'Page Summarization via Voice',
    'Voice Navigation On, any article/news page',
    '1. Navigate to a news article or blog post\n'
    '2. Say "summarize"\n'
    '3. Observe the summary panel',
    'Summary panel slides in from right side.\n'
    'Shows bullet points extracted from page content.\n'
    'Displays: tier (local), latency in ms.\n'
    'Buttons: Simplify, Read Aloud, Copy.',
    category='AI Engine')

add_tc('TC-035', 'Summary — Read Aloud',
    'Summary panel visible from TC-034',
    '1. Click "Read Aloud" button in summary panel\n'
    '2. Listen to the audio',
    'Browser speaks the summary text using Web Speech API.\n'
    'Speech rate is ~0.9x (slightly slower for clarity).\n'
    'Audio stops if page is navigated away.',
    category='AI Engine')

add_tc('TC-036', 'Summary — Copy',
    'Summary panel visible from TC-034',
    '1. Click "Copy" button in summary panel\n'
    '2. Paste into a text editor (Ctrl+V)',
    'Summary text is copied to clipboard.\n'
    'Pasted text matches the bullet points shown in panel.',
    category='AI Engine')

add_tc('TC-037', 'Email Summarization UI — Gmail',
    'Logged into Gmail, viewing an email thread',
    '1. Open Gmail and select an email\n'
    '2. Look for AccessBridge toolbar buttons above the email\n'
    '3. Click "Summarize" button\n'
    '4. Observe the summary panel',
    'Summarize and Simplify buttons injected into Gmail toolbar.\n'
    'Clicking Summarize extracts email body HTML.\n'
    'Summary panel shows email summary with bullet points.\n'
    'Reading time and complexity score displayed.',
    category='AI Engine')

add_tc('TC-038', 'Text Simplification via AI Bridge',
    'Any text-heavy page open',
    '1. Say "simplify" (voice command) or enable via Cognitive tab\n'
    '2. Observe the simplified overlay',
    'Simplified content overlay appears.\n'
    'Complex words replaced with simpler alternatives.\n'
    'Close button (×) dismisses the overlay.\n'
    'Original page content preserved underneath.',
    category='AI Engine')

add_tc('TC-039', 'AI Readability Score',
    'Any page open, service worker DevTools open',
    '1. In service worker console, run:\n'
    '   chrome.runtime.sendMessage({type:"AI_READABILITY",payload:{text:"This is simple."}})\n'
    '2. Check the response',
    'Returns {score: <number>, grade: "Easy"|"Medium"|"Hard"|"Very Hard"}.\n'
    'Score ≤6 = Easy, ≤10 = Medium, ≤14 = Hard, >14 = Very Hard.',
    category='AI Engine', priority='Medium')

# ============================================================
# SECTION 8: DOMAIN CONNECTORS
# ============================================================
doc.add_heading('8. Domain Connectors', level=1)

add_tc('TC-040', 'Banking Connector — Detection',
    'Extension loaded, banking website open (e.g. SBI/HDFC/ICICI)',
    '1. Navigate to any online banking website\n'
    '2. Open browser DevTools Console\n'
    '3. Look for "[AccessBridge] Domain detected: Banking"',
    'Banking connector auto-detects the domain.\n'
    'Console logs domain detection message.\n'
    'Jargon decoder activates (25 banking terms).\n'
    'Indian numbering: amounts shown in Lakh/Crore format.',
    category='Domain Connectors', priority='Medium')

add_tc('TC-041', 'Banking Connector — Jargon Decoder',
    'Banking website with financial terms visible',
    '1. Look for banking jargon on the page (e.g. "APR", "EMI", "CIBIL")\n'
    '2. Hover over or focus on jargon terms',
    'Jargon terms are decoded with simpler explanations.\n'
    'Tooltip or inline expansion shows plain-language meaning.\n'
    '25 banking terms supported.',
    category='Domain Connectors', priority='Medium')

add_tc('TC-042', 'Insurance Connector — Detection',
    'Extension loaded, insurance website open',
    '1. Navigate to an insurance website (e.g. LIC, ICICI Prudential)\n'
    '2. Check DevTools console for detection message',
    'Insurance connector auto-detects.\n'
    'Console: "[AccessBridge] Domain detected: Insurance".\n'
    'Policy simplifier and jargon decoder (35 terms) activated.',
    category='Domain Connectors', priority='Medium')

add_tc('TC-043', 'Domain Registry — No Match',
    'Extension loaded, generic website (e.g. google.com)',
    '1. Navigate to google.com\n'
    '2. Check DevTools console',
    'Console: "[AccessBridge] No domain-specific connector matched".\n'
    'No domain connector active — generic mode.\n'
    'All other features still work normally.',
    category='Domain Connectors', priority='Low')

# ============================================================
# SECTION 9: SIDE PANEL
# ============================================================
doc.add_heading('9. Side Panel', level=1)

add_tc('TC-044', 'Open Side Panel',
    'Extension loaded',
    '1. Right-click AccessBridge icon in toolbar\n'
    '2. Select "Open Side Panel"\n'
    '   OR: chrome://extensions > AccessBridge > Side Panel',
    'Side panel opens on the right side of browser.\n'
    'Shows AccessBridge dashboard UI.\n'
    'Rich content with struggle score gauge.',
    category='Side Panel')

add_tc('TC-045', 'Side Panel — Struggle Score Gauge',
    'Side panel open per TC-044',
    '1. Observe the struggle score display\n'
    '2. Browse and interact with pages\n'
    '3. Watch the score update',
    'Circular gauge shows current struggle score.\n'
    'Session timer shows elapsed time.\n'
    'Score updates in real-time.',
    category='Side Panel')

add_tc('TC-046', 'Side Panel — Adaptation History',
    'Side panel open, some features toggled on/off',
    '1. Enable Focus Mode from popup\n'
    '2. Disable Focus Mode\n'
    '3. Check adaptation history in side panel',
    'History log shows adaptations applied and reverted.\n'
    'Each entry has timestamp and type.',
    category='Side Panel')

add_tc('TC-047', 'Side Panel — Quick Controls',
    'Side panel open',
    '1. Locate the quick control grid (6 feature buttons)\n'
    '2. Click a control to toggle a feature',
    'Quick control grid shows 6 feature toggle buttons.\n'
    'Clicking toggles the feature on/off.\n'
    'Visual state updates to reflect active/inactive.',
    category='Side Panel')

# ============================================================
# SECTION 10: SELF-UPDATE
# ============================================================
doc.add_heading('10. Self-Update System', level=1)

add_tc('TC-048', 'No Update Banner When Current',
    'Extension v0.1.1 loaded, VPS API reports v0.1.1',
    '1. Open popup\n'
    '2. Observe the header area',
    'No update banner displayed.\n'
    'Header shows "AccessBridge v0.1.1".\n'
    'Version is read dynamically from manifest.',
    category='Self-Update')

add_tc('TC-049', 'Update Banner Appears When Outdated',
    'Extension v0.1.1, VPS API updated to report v0.1.2\n'
    '(Change CURRENT_VERSION in VPS API and restart)',
    '1. SSH to VPS, edit /opt/accessbridge/api/main.py\n'
    '   Change CURRENT_VERSION = "0.1.2"\n'
    '2. Run: docker compose restart accessbridge-api\n'
    '3. Open extension popup\n'
    '4. Observe update banner',
    'Purple gradient banner appears at top of popup.\n'
    'Shows "v0.1.2 available" with changelog text.\n'
    'Two buttons: "Download" and "Reload".\n'
    'Download button opens zip URL in new tab.',
    category='Self-Update')

add_tc('TC-050', 'Reload Button Restarts Extension',
    'Update banner visible per TC-049',
    '1. Click "Reload" button in update banner\n'
    '2. Observe extension behavior',
    'Extension calls chrome.runtime.reload().\n'
    'Popup closes (extension restarts).\n'
    'After restart, if dist/ was updated, new version loads.\n'
    'If dist/ not updated, same version reloads (banner reappears).',
    category='Self-Update')

# ============================================================
# SECTION 11: PROFILE MANAGEMENT
# ============================================================
doc.add_heading('11. Profile Management', level=1)

add_tc('TC-051', 'Profile Persists Across Sessions',
    'Extension loaded with default profile',
    '1. Open popup > Sensory tab\n'
    '2. Set Font Scale to 1.5x\n'
    '3. Set Contrast to 1.3x\n'
    '4. Close popup\n'
    '5. Reopen popup > Sensory tab',
    'Font Scale still shows 1.5x.\n'
    'Contrast still shows 1.3x.\n'
    'Profile saved to chrome.storage.local and persists.',
    category='Profile')

add_tc('TC-052', 'Export Profile',
    'Popup open, Settings tab, some settings changed',
    '1. Go to Settings tab\n'
    '2. Click "Export Profile"\n'
    '3. Check downloaded file',
    'JSON file downloads: "accessbridge-profile.json".\n'
    'File contains full profile with all settings.\n'
    'JSON is valid and human-readable.',
    category='Profile')

add_tc('TC-053', 'Import Profile',
    'Popup open, Settings tab, a previously exported JSON file',
    '1. Go to Settings tab\n'
    '2. Click "Import Profile"\n'
    '3. Select a previously exported JSON file\n'
    '4. Check that settings update',
    'File picker opens, accepts .json files.\n'
    'Profile loads and all settings update immediately.\n'
    'Switching to Sensory/Cognitive/Motor tabs shows imported values.',
    category='Profile')

add_tc('TC-054', 'Language Setting',
    'Popup open, Settings tab',
    '1. Change Language dropdown to "Hindi"\n'
    '2. Close and reopen popup\n'
    '3. Check Settings tab',
    'Language persists as "Hindi" after reopen.\n'
    'Voice commands switch to hi-IN recognition.\n'
    '8 languages available: English, Spanish, French, German, Chinese, Japanese, Arabic, Hindi.',
    category='Profile')

# ============================================================
# SECTION 12: VPS & LANDING PAGE
# ============================================================
doc.add_heading('12. VPS & Landing Page', level=1)

add_tc('TC-055', 'Landing Page Loads',
    'Internet access available',
    '1. Navigate to http://72.61.227.64:8300/\n'
    '2. Observe the page',
    'Landing page loads with full-width layout.\n'
    'Navbar: AccessBridge logo, How It Works, Features, Architecture, GitHub, Install button.\n'
    'Hero section with title and description.\n'
    'Stats cards, features grid, architecture section.\n'
    'Footer shows dynamic version and "Team: Manish Kumar".',
    category='VPS')

add_tc('TC-056', 'Landing Page — Dynamic Version',
    'Landing page loaded per TC-055',
    '1. Scroll to footer\n'
    '2. Check version number\n'
    '3. Open DevTools > Network tab\n'
    '4. Refresh page and find /api/version request',
    'Footer shows version fetched from /api/version endpoint.\n'
    'Network tab shows successful GET to /api/version.\n'
    'Response JSON: {version, download_url, changelog}.\n'
    'Health dot is green (API reachable).',
    category='VPS')

add_tc('TC-057', 'Extension Download from Landing Page',
    'Landing page loaded',
    '1. Click "Download Extension (.zip)" button\n'
    '2. Save the file\n'
    '3. Unzip and inspect contents',
    'ZIP file downloads (~135KB).\n'
    'Contains: manifest.json, src/, assets/, icons/.\n'
    'manifest.json shows version 0.1.1.\n'
    'Can be loaded as unpacked extension in Chrome.',
    category='VPS')

add_tc('TC-058', 'Landing Page — Mobile Responsive',
    'Landing page loaded on mobile or DevTools mobile emulation',
    '1. Open DevTools > Toggle Device Toolbar\n'
    '2. Select iPhone 14 Pro or similar\n'
    '3. Observe layout',
    'Page fills full viewport width — no horizontal scroll.\n'
    'Navbar collapses to hamburger menu.\n'
    'Stats grid: 2 columns.\n'
    'Features grid: single column.\n'
    'Install steps stack vertically.\n'
    'All text readable, buttons full-width.',
    category='VPS')

# ============================================================
# SECTION 13: REGRESSION & EDGE CASES
# ============================================================
doc.add_heading('13. Regression & Edge Cases', level=1)

add_tc('TC-059', 'Multiple Tabs — Adaptations Per Tab',
    'Extension loaded, two tabs open with different sites',
    '1. Open Tab A (e.g. Wikipedia)\n'
    '2. Enable Focus Mode\n'
    '3. Open Tab B (e.g. Google)\n'
    '4. Check if Focus Mode is active on Tab B',
    'Focus Mode applies to both tabs (content script runs on all).\n'
    'Toggling off in popup affects the active tab.\n'
    'REVERT_ALL message sent to all tabs.',
    category='Regression')

add_tc('TC-060', 'Extension Survives Page Navigation',
    'Extension loaded, Focus Mode enabled',
    '1. Enable Focus Mode on a page\n'
    '2. Click a link to navigate to another page\n'
    '3. Check if content script re-initializes',
    'Content script re-initializes on new page.\n'
    'Console: "[AccessBridge] Detected app: generic".\n'
    'Profile-based features auto-start if enabled.\n'
    'Manual toggles need re-activation.',
    category='Regression')

add_tc('TC-061', 'No Console Errors on chrome:// Pages',
    'Extension loaded',
    '1. Navigate to chrome://settings\n'
    '2. Open DevTools console\n'
    '3. Check for AccessBridge errors',
    'Content scripts do not inject on chrome:// pages.\n'
    'No errors related to AccessBridge.\n'
    'Extension popup still works normally.',
    category='Regression')

add_tc('TC-062', 'Concurrent Features — No Conflicts',
    'Extension loaded, any webpage',
    '1. Enable Focus Mode + Voice Navigation + Dwell Click simultaneously\n'
    '2. Interact with the page\n'
    '3. Use voice to navigate while dwell click is active',
    'All three features work simultaneously.\n'
    'No JavaScript errors.\n'
    'Focus spotlight, voice listening, and dwell click all functional.\n'
    'Performance remains acceptable (no visible lag).',
    category='Regression')

add_tc('TC-063', 'REVERT_ALL Cleans Up Everything',
    'Multiple features enabled (Focus Mode, Voice Nav, Dwell Click, Keyboard Mode)',
    '1. Open popup\n'
    '2. Toggle master switch Off\n'
    '3. Inspect the page for any remaining AccessBridge UI elements',
    'All overlays, cursors, skip links, prediction panels removed.\n'
    'Voice recognition stops.\n'
    'Dwell click stops.\n'
    'Eye tracker stops and webcam released.\n'
    'Page returns to original state with no AccessBridge artifacts.',
    category='Regression')


# ---- Summary Table ----
doc.add_page_break()
doc.add_heading('Test Summary', level=1)

summary_table = doc.add_table(rows=1, cols=4, style='Table Grid')
summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, header in enumerate(['Category', 'Test Cases', 'Priority', 'Count']):
    cell = summary_table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    run.bold = True
    run.font.size = Pt(9)

categories = [
    ('Installation', 'TC-001 to TC-003', 'High', '3'),
    ('Popup UI', 'TC-004 to TC-009', 'High', '6'),
    ('Sensory', 'TC-010 to TC-016', 'High', '7'),
    ('Cognitive', 'TC-017 to TC-023', 'High', '7'),
    ('Motor', 'TC-024 to TC-033', 'High', '10'),
    ('AI Engine', 'TC-034 to TC-039', 'High', '6'),
    ('Domain Connectors', 'TC-040 to TC-043', 'Medium', '4'),
    ('Side Panel', 'TC-044 to TC-047', 'Medium', '4'),
    ('Self-Update', 'TC-048 to TC-050', 'High', '3'),
    ('Profile', 'TC-051 to TC-054', 'High', '4'),
    ('VPS & Landing Page', 'TC-055 to TC-058', 'Medium', '4'),
    ('Regression', 'TC-059 to TC-063', 'High', '5'),
    ('TOTAL', '', '', '63'),
]

for cat, tests, pri, count in categories:
    row = summary_table.add_row()
    for i, val in enumerate([cat, tests, pri, count]):
        run = row.cells[i].paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        if cat == 'TOTAL':
            run.bold = True

# Save
output_path = r'E:\code\AccessBridge\AccessBridge_Manual_Test_Cases.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Total test cases: {tc_counter[0]}')
